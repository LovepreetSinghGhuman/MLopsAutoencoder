from docx import Document

doc = Document()
doc.add_heading('MLOps Fraud Detection with Autoencoders – Assignment Report', 0)

doc.add_heading('GitHub Repository', level=1)
doc.add_paragraph('• MLopsAutoencoder GitHub Repository: https://github.com/LovepreetSinghGhuman/MLopsAutoencoder.git')
doc.add_paragraph('• Demo Video: https://drive.google.com/file/d/18_wjC4EDZMhxeHaHiFTGqlicmHjhTNaA/view?usp=drive_link')

doc.add_heading('Project Overview', level=1)
doc.add_paragraph(
    "This project demonstrates a complete MLOps workflow for fraud detection using an autoencoder neural network. "
    "The solution covers data preprocessing, model training on Azure ML, automated CI/CD with GitHub Actions, and deployment to Kubernetes (AKS) with a user-friendly frontend."
)

doc.add_heading('Data & Preprocessing', level=2)
doc.add_paragraph(
    "- Dataset: IEEE-CIS Fraud Detection dataset (preprocessed).\n"
    "- Preprocessing: Data cleaning, feature engineering, and scaling are performed in both the notebook and src/train.py. "
    "Steps include log-transforming transaction amounts, device info grouping, and feature scaling.\n"
    "- Artifacts: Cleaned and reduced datasets are stored in Data/Processed/."
)

doc.add_heading('Model', level=2)
doc.add_paragraph(
    "- Type: Autoencoder neural network for anomaly detection.\n"
    "- Training Script: src/train.py\n"
    "- Artifacts: Model, scaler, config, and threshold are saved in models/."
)

doc.add_heading('Task 1: Cloud Training (Azure ML)', level=1)
doc.add_paragraph(
    "- Service: Azure ML is used for model training and retraining.\n"
    "- Automation: Training is triggered by GitHub Actions (.github/workflows/azure-pipelines-ci-cd.yaml), which submits a pipeline job (deployment/pipeline-job.yaml) using a registered component (deployment/train-component.yaml).\n"
    "- Artifacts: Model files are downloaded after training and used for deployment.\n"
    "- Reproducibility: All preprocessing steps are consistent between notebook, training script, and API.\n"
    "Screenshots to include:\n"
    "- Azure ML job submission and completion (portal or CLI).\n"
    "- Pipeline run in GitHub Actions."
)

doc.add_heading('Task 2: Kubernetes Deployment', level=1)
doc.add_paragraph(
    "- Backend: FastAPI app (src/score.py), containerized with deployment/Dockerfile.\n"
    "- Frontend: Responsive HTML/CSS/JS (frontend/index.html, frontend/style.css, frontend/script.js), served by NGINX (frontend/Dockerfile).\n"
    "- Kubernetes Manifests: Backend (deployment/k8s/deployment.yaml, deployment/k8s/service.yaml), Frontend (frontend/frontend-nginx-deployment.yaml, frontend/frontend-nginx-service.yaml), Ingress (deployment/k8s/ingress.yaml).\n"
    "- Ingress: NGINX Ingress controller is installed and configured for routing and public access.\n"
    "- User Interaction: Users upload a CSV/Excel file and view predictions in a scrollable, styled table.\n"
    "- Microservice Communication: [User] --> [Frontend (NGINX)] --> [K8s Ingress] --> [FastAPI (Autoencoder)] --> [Model Artifacts]\n"
    "Screenshots to include:\n"
    "- Frontend UI (file upload and result table).\n"
    "- FastAPI /docs page.\n"
    "- Output of kubectl get pods,svc,ingress."
)

doc.add_heading('Task 3: CI/CD Automation (GitHub Actions)', level=1)
doc.add_paragraph(
    "- Retraining: On every push to main, the workflow triggers a new Azure ML training job.\n"
    "- Redeployment: After training, the workflow builds and pushes new Docker images for backend and frontend, and updates the AKS deployments.\n"
    "- Minimal Manual Steps: All steps are automated after code is pushed to main.\n"
    "- Model Versioning: Handled via Azure ML's model registry.\n"
    "- Secrets: Managed via GitHub Actions secrets.\n"
    "Screenshots to include:\n"
    "- GitHub Actions workflow runs (showing retrain, build, deploy steps)."
)

doc.add_heading('Reflection', level=1)
doc.add_paragraph(
    "During this project, I learned how to orchestrate a full MLOps pipeline using Azure ML, Kubernetes, and GitHub Actions. "
    "The most challenging part was ensuring reproducibility between local and cloud environments, especially for data preprocessing and model artifact management. "
    "I solved this by centralizing all preprocessing logic in shared scripts and automating artifact handling in the pipeline. "
    "Setting up the NGINX Ingress controller and ARM64 Docker builds for AKS also required extra troubleshooting, but provided valuable experience with real-world cloud deployment issues."
)

doc.add_heading('Demo Video', level=1)
doc.add_paragraph(
    "A short demo video is included, showing:\n"
    "- The frontend UI for file upload and result display.\n"
    "- The FastAPI /docs page.\n"
    "- The running Kubernetes cluster (kubectl get pods,svc,ingress)."
)

doc.add_heading('Additional Notes', level=1)
doc.add_paragraph(
    "- All files are UTF-8 encoded and checked in the pipeline.\n"
    "- Model artifacts are managed by the pipeline and not tracked in git.\n"
    "- For any issues, logs are available in GitHub Actions and Kubernetes pods."
)

doc.add_paragraph("\nAuthor:\nLovepreet Singh\nMLOps and AI design patterns")

doc.save("MLopsAutoencoder_Report.docx")